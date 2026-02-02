import os
import glob
from pathlib import Path
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.embeddings import Embeddings  # ✅ ДОБАВЛЕНО
from langchain_openai import ChatOpenAI
from sentence_transformers import SentenceTransformer
import shutil

# Unstructured (опционально)
try:
    from unstructured.partition.docx import partition_docx
    from unstructured.partition.pdf import partition_pdf
    UNSTRUCTURED_AVAILABLE = True
    print("✅ Unstructured доступен")
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    print("⚠️  Установите: pip install 'unstructured[docx,pdf]' для DOCX/PDF")

class LocalEmbeddings(Embeddings):  # ✅ НАСЛЕДУЕМ Embeddings!
    def __init__(self, model_name="paraphrase-multilingual-mpnet-base-v2"):
        print("🔄 Загрузка эмбеддингов (multilingual для RU)...")
        self.model = SentenceTransformer(model_name)
        print("✅ Эмбеддинги готовы")
    
    def embed_documents(self, texts):
        return self.model.encode(texts).tolist()
    
    def embed_query(self, text):
        return self.model.encode([text])[0].tolist()

class RobustDocumentLoader:
    @staticmethod
    def load_txt(file_path: str) -> List[Document]:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read().strip()
                if content:
                    return [Document(page_content=content, metadata={"source": file_path})]
        except Exception as e:
            print(f"⚠️  TXT {file_path}: {e}")
        return []
    
    @staticmethod
    def load_docx(file_path: str) -> List[Document]:
        # Игнор временных файлов
        if file_path.endswith('~$'):
            return []
            
        if UNSTRUCTURED_AVAILABLE:
            try:
                elements = partition_docx(file_path)
                text = '\n'.join([el.text for el in elements if el.text and el.text.strip()])
                if text.strip():
                    return [Document(page_content=text.strip(), metadata={"source": file_path})]
            except Exception as e:
                print(f"⚠️  Unstructured DOCX {file_path}: {e}")
        
        # Fallback python-docx ✅ ФИКС
        try:
            from docx import Document
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            if text.strip():
                return [Document(page_content=text.strip(), metadata={"source": file_path})]
        except Exception as e:
            print(f"⚠️  DOCX {file_path}: {e}")
        return []
    
    @staticmethod
    def load_pdf(file_path: str) -> List[Document]:
        if UNSTRUCTURED_AVAILABLE:
            try:
                elements = partition_pdf(file_path)
                text = '\n'.join([el.text for el in elements if el.text and el.text.strip()])
                if text.strip():
                    return [Document(page_content=text.strip(), metadata={"source": file_path})]
            except Exception as e:
                print(f"⚠️  Unstructured PDF {file_path}: {e}")
        
        # Fallback PyPDF2
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                return [Document(page_content=text.strip(), metadata={"source": file_path})]
        except Exception as e:
            print(f"⚠️  PDF {file_path}: {e}")
        return []

def load_knowledge_base(folder: str = "./knowledge_base/") -> List[Document]:
    if not Path(folder).exists():
        print(f"❌ Папка {folder} не найдена!")
        return []
        
    print("📂 Сканирование knowledge_base/")
    all_docs = []
    
    # TXT
    for file in Path(folder).rglob("*.txt"):
        docs = RobustDocumentLoader.load_txt(str(file))
        all_docs.extend(docs)
        if docs:
            print(f"📄 TXT: {file.name}")
    
    # DOCX
    for file in Path(folder).rglob("*.docx"):
        docs = RobustDocumentLoader.load_docx(str(file))
        all_docs.extend(docs)
        if docs:
            print(f"📄 DOCX: {file.name}")
    
    # PDF
    for file in Path(folder).rglob("*.pdf"):
        docs = RobustDocumentLoader.load_pdf(str(file))
        all_docs.extend(docs)
        if docs:
            print(f"📄 PDF: {file.name}")
    
    print(f"\n✅ Загружено {len(all_docs)} документов")
    return all_docs

class RAGPro:
    def __init__(self):
        with open("api_key.txt", "r", encoding='utf-8') as f:
            self.api_key = f.read().strip()
        
        self.embeddings = LocalEmbeddings()
        self.llm = ChatOpenAI(model="gpt-4o-mini", api_key=self.api_key, temperature=0)
        self.vectorstore = None
        self.chain = None
        
    def build_index(self, docs: List[Document]):
        if os.path.exists("faiss_index"):
            shutil.rmtree("faiss_index")
            print("🗑️  Удалён старый faiss_index")
        
        if not docs:
            print("❌ Нет документов для индекса!")
            return None
            
        splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=300)
        splits = splitter.split_documents(docs)
        print(f"✂️  Разбито на {len(splits)} чанков")
        
        self.vectorstore = FAISS.from_documents(splits, self.embeddings)
        self.vectorstore.save_local("faiss_index")
        
        retriever = self.vectorstore.as_retriever(search_kwargs={"k": 20})
        
        template = """Ты эксперт-консультант по базе знаний. Используй ТОЛЬКО информацию из КОНТЕКСТА.

ПРАВИЛА:
1. Отвечай точно, цитируя источник (файл).
2. Нет данных → "Нет информации в базе знаний".
3. Русский язык, кратко.

КОНТЕКСТ:
{context}

ВОПРОС: {question}

ОТВЕТ:"""
        
        prompt = ChatPromptTemplate.from_template(template)
        
        def format_context(docs):
            return "\n\n".join(f"Источник: {doc.metadata.get('source', 'Неизвестно')}\n{doc.page_content}" for doc in docs)
        
        chain = (
            {"context": retriever | format_context, "question": RunnablePassthrough()}
            | prompt
            | self.llm
            | StrOutputParser()
        )
        self.chain = chain
        
        print(f"✅ FAISS готов (k=20, {len(splits)} чанков)")
        return self.vectorstore
    
    def ask(self, question: str) -> str:
        return self.chain.invoke(question)
    
    def extract(self, question: str) -> str:
        retriever = self.vectorstore.as_retriever(search_kwargs={"k": 25})
        docs = retriever.invoke(question)
        context = "\n\n".join(doc.page_content for doc in docs)
        prompt = f"""ИЗВЛЕКИ точный ответ ИЗ КОНТЕКСТА. Цитируй. Нет → "НЕ НАШЁЛ".

КОНТЕКСТ: {context}

ВОПРОС: {question}

ОТВЕТ:"""
        return self.llm.invoke(prompt).content.strip()

def main():
    print("🚀 RAG Pro v3.1 - Фикс Windows/Python 3.14")
    
    docs = load_knowledge_base()
    if not docs:
        print("❌ Добавьте файлы в knowledge_base/ и перезапустите")
        return
    
    rag = RAGPro()
    rag.build_index(docs)
    
    print("\n🧪 ТЕСТ:")
    tests = ["кто директор?", "стоимость контракта", "пункт 2.3"]
    for test in tests:
        answer = rag.ask(test)
        print(f"❓ {test}\n💬 {answer}\n")
    
    print("🤖 ЧАТ (exit/выход | ! для extract):")
    while True:
        try:
            question = input("\n❓ ").strip()
            if question.lower() in ['exit', 'выход', 'quit']:
                break
            if question.startswith("!"):
                answer = rag.extract(question[1:])
            else:
                answer = rag.ask(question)
            print(f"💬 {answer}")
        except KeyboardInterrupt:
            print("\n👋")
            break

if __name__ == "__main__":
    main()
